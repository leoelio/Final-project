from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "glasgow_msc_thesis_and_poster_execution_guide_zh.docx"

NAVY = "1F4D78"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E8F2EE"
PALE_RED = "F8ECEC"
INK = "1F2933"
MUTED = "5B6573"
GRID = "C8D3E0"

USABLE_WIDTH = 9360


def set_run_font(run, size=11, bold=None, color=INK, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    """Keep tblW, tblGrid and every tcW in one DXA coordinate system."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for idx, width in enumerate(widths):
        grid.gridCol_lst[idx].set(qn("w:w"), str(width))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_borders(table, color=GRID):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_bottom_border(paragraph, color=NAVY, size="16"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "8")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    set_run_font(run, size=8, color=MUTED)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 25, NAVY, 0, 6),
        ("Subtitle", 12, MUTED, 0, 12),
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, NAVY, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = name.startswith("Heading")

    if "Guide Note" not in document.styles:
        note = document.styles.add_style("Guide Note", WD_STYLE_TYPE.PARAGRAPH)
        note.base_style = normal
        note.font.name = "Calibri"
        note._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        note.font.size = Pt(10)
        note.font.color.rgb = RGBColor.from_string(INK)
        note.paragraph_format.space_before = Pt(3)
        note.paragraph_format.space_after = Pt(8)
        note.paragraph_format.left_indent = Inches(0.12)
        note.paragraph_format.right_indent = Inches(0.12)
        note.paragraph_format.line_spacing = 1.15

    bullet = document.styles["List Bullet"]
    bullet.font.name = "Calibri"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    bullet.font.size = Pt(10.5)
    bullet.font.color.rgb = RGBColor.from_string(INK)
    bullet.paragraph_format.left_indent = Inches(0.32)
    bullet.paragraph_format.first_line_indent = Inches(-0.16)
    bullet.paragraph_format.space_after = Pt(3)
    bullet.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("MUJOCO WIDOWX | MSc THESIS + POSTER EXECUTION GUIDE")
    set_run_font(header_run, size=8, bold=True, color=NAVY)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Research-freeze writing guide | Page ")
    set_run_font(footer_run, size=8, color=MUTED)
    add_page_field(footer)


def add_title(document, text, subtitle=None):
    paragraph = document.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=25, bold=True, color=NAVY)
    add_bottom_border(paragraph)
    if subtitle:
        sub = document.add_paragraph(style="Subtitle")
        run = sub.add_run(subtitle)
        set_run_font(run, size=12, color=MUTED)


def add_text(document, text, bold_prefix=None):
    paragraph = document.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_note(document, label, text, fill=PALE_BLUE):
    paragraph = document.add_paragraph(style="Guide Note")
    set_paragraph_shading(paragraph, fill)
    label_run = paragraph.add_run(label + "  ")
    set_run_font(label_run, size=10, bold=True, color=NAVY if fill != PALE_RED else "8C2D2D")
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10, color=INK)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    return paragraph


def add_table(document, headers, rows, widths, compact=False):
    table = document.add_table(rows=1, cols=len(headers))
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        set_run_font(run, size=9 if compact else 9.5, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            if isinstance(value, tuple):
                value, fill = value
                set_cell_shading(cell, fill)
            if isinstance(value, str) and value.startswith("[OK]"):
                set_cell_shading(cell, PALE_GREEN)
                value = value.replace("[OK]", "", 1).strip()
            elif isinstance(value, str) and value.startswith("[NO]"):
                set_cell_shading(cell, PALE_RED)
                value = value.replace("[NO]", "", 1).strip()
            run = paragraph.add_run(str(value))
            set_run_font(run, size=8.8 if compact else 9.3)
    set_table_geometry(table, widths)
    set_table_borders(table)
    document.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_path(document, path):
    paragraph = document.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("Evidence source: " + path)
    set_run_font(run, size=9, italic=True, color=MUTED)


def build_document():
    document = Document()
    configure_document(document)

    add_title(
        document,
        "格拉斯哥大学工学院 MSc 论文与海报执行建议",
        "基于已冻结的 MuJoCo WidowX 研究证据包 | 中文工作指南 | 2026-08-03",
    )
    add_text(
        document,
        "用途：把目前已经完成、可复核的 MuJoCo 研究整理成一篇诚实、逻辑完整的 MSc dissertation 和一张可答辩的 poster。本文不新增研究，不补写未经验证的结果，也不把后续计划包装成完成成果。",
    )
    add_note(
        document,
        "核心叙事",
        "项目的价值不在于声称已经完成 OpenVLA LoRA 或真实机械臂迁移，而在于：在有限算力、有限示范和 MuJoCo-only 条件下，系统性比较了轻量模仿学习与 PEFT/VLM proxy，最终形成一个可审计的“冻结 CLIP 语义 + 顶视 RGB 几何 + 结构化执行 + 一次有界重试”的模块化视觉-语言-动作基线。",
    )
    add_table(
        document,
        ["当前状态", "可在论文中如何使用"],
        [
            ["已完成 MuJoCo 研究包", "可作为 dissertation 的方法、评测与复现实验主体。"],
            ["最终 V4 两组独立复核", "主结果：严格任务成功 278/288（描述性合并），但不得写成对所有学习方法的因果优越性。"],
            ["CLIP LoRA spatial pointer", "可作为轻量后训练的负结果：3/20 闭环成功，未通过部署门槛。"],
            ["OpenVLA / Isaac / real WidowX", "只有桥接、可行性或交接材料；只能写作 future work，不可写成实验结果。"],
        ],
        [2500, 6860],
    )
    add_text(document, "制作前的唯一原则：每一条性能结论必须同时能追溯到原始评测记录；视频只用于展示执行过程，不替代 aggregate 统计。", bold_prefix="制作前的唯一原则：")

    document.add_page_break()
    document.add_heading("1. 先把研究讲成一条完整的科学故事", level=1)
    add_text(
        document,
        "不要把这篇论文写成“做了很多方法”的流水账。建议采用“问题 - 证据 - 决策”的叙事：先说明资源受限条件下直接学习连续机械臂动作的困难，再展示小型 BC、trajectory/ACT/diffusion 以及 action-head / LoRA-style proxy 的闭环瓶颈，最后解释为何转向职责分解的 V4，并以独立 seed 复核和负向干预试验结束。",
    )
    add_table(
        document,
        ["阶段", "已经得到的材料", "论文中承担的作用"],
        [
            ["任务与数据链", "WidowX MuJoCo 场景、scripted expert、演示采集、.npz 回放、固定 seed 评测、viewer/MP4 导出。", "证明环境、数据与复现链路是可运行的。"],
            ["普通 IL 对照", "Linear BC、MLP BC、kNN BC；均未形成稳定闭环抓放。", "建立问题难度与简单模仿学习的基线。"],
            ["时序与生成式对照", "Trajectory-conditioned BC、ACT-style、Diffusion-style 的本地轻量实现。", "说明仅增加历史或动作块并未解决接触和抬升不稳定。"],
            ["轻量后训练探索", "Action head、Adapter、LoRA-style proxy、frozen CLIP pointer、CLIP visual LoRA pointer。", "把“PEFT 是否足够”作为可检验的失败假设，而非宣传性成果。"],
            ["最终可部署基线", "V4：冻结 CLIP 闭集语义、RGB 对象/目标定位、PickPlaceExpert、最多一次 RGB 重定位重试。", "回答在当前资源条件下什么方案最稳健、最可复核。"],
            ["审计与边界", "两组 seed-disjoint V4 复核、contact monitor 反例、counterfactual intervention audit、证据门禁。", "把负结果与限制作成贡献，而不是藏起来。"],
        ],
        [1450, 4550, 3360],
        compact=True,
    )
    add_path(document, "docs/claim_evidence_traceability.md; docs/method_evidence_gate.md; docs/final_experiment_package.md")
    add_note(
        document,
        "写作重点",
        "把“最终方案为何不是学习动作头”说清楚：运行时由语言和 RGB 决策目标，连续抓取/放置轨迹由结构化执行器负责。MuJoCo 物体真值只用于离线评分，不参与对象选择、重试判断或轨迹规划。",
        fill=PALE_GREEN,
    )

    document.add_heading("2. 可以写、必须谨慎写、不能写的边界", level=1)
    add_table(
        document,
        ["类别", "可用表述", "禁用或需改写的表述"],
        [
            ["最终 V4", "“在两个独立 MuJoCo seed cohort 上可重复执行指定桌面任务。”", "“端到端 VLA 已超过所有学习方法。”"],
            ["LoRA", "“完成了 CLIP 视觉 backbone 最后两层的低秩适配探索，未通过闭环部署门槛。”", "“完成 OpenVLA LoRA 微调”或“VLA 微调成功”。"],
            ["BC / ACT / Diffusion", "“本地轻量 baseline 在当前闭环接触任务中不稳定。”", "“复现了完整官方 visual ACT / Diffusion Policy”。"],
            ["语言能力", "“对预定义的闭集指令执行语义选择；保留了语言/空间测试。”", "“具备开放式自然语言泛化能力”。"],
            ["仿真到真实", "“完成 MuJoCo 扰动代理、真实机和 OpenVLA 的交接与数据桥接准备。”", "“已完成 Isaac 或真实 WidowX 迁移验证”。"],
            ["视频", "“视频用于定性复核，数值结论以固定 seed 的 aggregate JSON/CSV 为准。”", "“单条成功视频证明整体成功率”。"],
        ],
        [1500, 3900, 3960],
        compact=True,
    )

    document.add_heading("3. 论文定位、研究问题与贡献", level=1)
    add_text(document, "推荐英文论文题目：")
    add_note(
        document,
        "Recommended title",
        "Resource-Constrained Vision-Language-Action for MuJoCo WidowX Tabletop Manipulation: From Lightweight Adaptation to RGB-Grounded Structured Control.",
    )
    add_text(document, "推荐中文题目：资源受限条件下用于 MuJoCo WidowX 桌面操作的视觉-语言-动作研究：从轻量适配到 RGB 驱动的结构化控制。")
    add_text(document, "主研究问题建议收敛为：在有限算力、小规模 MuJoCo scripted demonstrations 且没有真实机械臂验证的条件下，哪些轻量化视觉-语言-动作适配路线能形成可靠的闭环桌面抓放策略？若直接学习路线失败，模块化语义、视觉几何和结构化控制的分解能否提供可重复的 MuJoCo 基线？")
    document.add_heading("建议保留的三项贡献", level=2)
    for item in [
        "构建了可复现实验闭环：MuJoCo WidowX 场景、任务定义、演示记录与回放、固定 seed 评测、严格评分、viewer 与视频证据索引。",
        "在统一的闭环任务与资源记录下，对普通 BC、时序/生成式轻量 baseline、action-head / PEFT proxy 以及 CLIP spatial adaptation 进行了系统性筛选，并报告失败边界。",
        "提出并独立复核 V4 模块化基线：冻结 CLIP 负责闭集任务语义，顶视 RGB 负责源/目标几何，结构化执行器负责接触运动，并通过一次有界恢复处理首轮失败。",
    ]:
        add_bullet(document, item)
    document.add_heading("推荐的研究问题层级", level=2)
    add_table(
        document,
        ["RQ", "论文回答方式", "不要扩大的范围"],
        [
            ["RQ1: 小数据模仿学习能否稳定闭环？", "用 BC、kNN、trajectory/ACT/diffusion local baseline 的严格闭环表现说明其限制。", "不把训练范围记忆当作泛化。"],
            ["RQ2: 轻量适配 / LoRA 是否解决空间 grounding？", "报告 frozen 与 LoRA CLIP pointer 的相同 holdout 和部署门槛。", "不将 CLIP LoRA 称作 OpenVLA LoRA。"],
            ["RQ3: 当前资源下的可用方案是什么？", "用 V4 两个独立 cohort、语义/对象选择/严格成功与首轮成功指标回答。", "不声称真实机器人或开放世界适用。"],
            ["RQ4: 为什么拒绝新的接触策略？", "以 offline accuracy 与闭环 regressions 的差异解释 contact monitor 被拒绝。", "不以离线分类分数取代反事实安全性。"],
        ],
        [2100, 4450, 2810],
        compact=True,
    )
    document.add_heading("建议的起草顺序", level=2)
    for item in [
        "先锁定 Results 的 V4 主表与 LoRA/contact monitor 负结果表，确保每一个数值都有唯一来源。",
        "再写 Environment、Protocol 和 Methods；让读者在看结果之前已理解任务、strict metric、seed 与运行时信息边界。",
        "最后写 Introduction、Discussion 和 Conclusion；这些章节应由证据反推，而不是先写结论再寻找支持。",
    ]:
        add_bullet(document, item)

    document.add_page_break()
    document.add_heading("4. 推荐的 dissertation 章节结构", level=1)
    add_text(document, "下列篇幅只是一份写作预算，不是学校规定。最终字数、封面、模板、引用格式、附录和提交方式必须以你的 programme handbook、supervisor 指示与入学年度规定为准。")
    add_table(
        document,
        ["章节", "建议内容与写作目标", "优先证据来源"],
        [
            ["1. Introduction", "资源约束、桌面操作难点、研究问题、三项贡献、边界声明。不要用大模型口号替代问题定义。", "final_research_closed_loop_v2.md; claim_evidence_traceability.md"],
            ["2. Related Work", "BC/ACT/Diffusion、VLM/VLA、PEFT/LoRA、视觉伺服与 sim-to-real。使用原始论文和官方仓库，不用本项目报告充当 related work。", "外部文献库；项目只提供问题脉络"],
            ["3. Environment and Protocol", "WidowX/MJCF、相机、动作、任务、物体、示范、严格成功定义、seed、数据切分、视频与复现规则。", "task_bc_stage_report.md; reproducible_command_index.md"],
            ["4. Methods", "方法族分类；V4 数据流；为何 state 只作评分；LoRA pointer 的低秩适配边界。完整方法表放 appendix。", "method_evidence_gate.md; clip_lora_patch_pointer_stage_v1.md"],
            ["5. Results", "先放 aggregate quantitative results，再放分任务/泛化/资源表；视频只作质性说明。", "final_closure_audit_v1.md; result_matrix.md; *.csv"],
            ["6. Analysis and Discussion", "失败模式、V4 的职责分解、contact monitor 的离线-闭环脱节、统计解释、external validity。", "contact_phase_monitor_heldout_v1_analysis.md; counterfactual audit"],
            ["7. Conclusion and Future Work", "直接回答四个 RQ；总结贡献；明确 OpenVLA、Isaac、真实机只属于未来工作。", "final_research_closed_loop_v2.md"],
            ["Appendices", "完整方法/参数/命令/更多表、证据路径、视频索引、任务样例。", "thesis_appendix_tables.md; repository_reproducibility.md"],
        ],
        [1850, 4880, 2630],
        compact=True,
    )

    document.add_page_break()
    document.add_heading("5. 结果章节应如何组织", level=1)
    add_note(
        document,
        "推荐顺序",
        "先定义严格成功与 cohort，再给主表；随后给普通 baseline 和 LoRA 负结果；最后解释为何 V4 被保留、为何 contact monitor 被拒绝。不要按脚本创建时间排列结果。",
    )
    document.add_heading("5.1 论文主表必须出现的 V4 结果", level=2)
    add_table(
        document,
        ["独立 cohort", "Seeds", "严格成功", "语义正确", "初始对象选择", "首轮成功", "Wilson 95% CI"],
        [
            ["rgb_table_recovery_v4_extended_v1", "4000-4011", "135/144 (93.8%)", "144/144", "144/144", "127/144", "0.885-0.967"],
            ["contact_phase_monitor_heldout_v1 / v4_standard", "10000-10011", "143/144 (99.3%)", "144/144", "144/144", "138/144", "0.962-0.999"],
            [("描述性合并", PALE_GREEN), "两个不重叠 cohort", ("278/288 (96.5%)", PALE_GREEN), "288/288", "288/288", "265/288", "0.937-0.981"],
        ],
        [2160, 1100, 1330, 1100, 1300, 1150, 1220],
        compact=True,
    )
    add_text(document, "主表下的推荐英文说明：The pooled value is descriptive evidence of repeatability across two seed-disjoint cohorts; it is not a causal comparison against another method.")
    add_path(document, "docs/final_closure_audit_v1.md; docs/v4_independent_replication_v1.md")

    document.add_heading("5.2 负结果不是附带内容，而是方法选择证据", level=2)
    add_table(
        document,
        ["候选", "关键记录", "论文结论"],
        [
            ["CLIP visual LoRA Patch Pointer", "LoRA 24,576 + pointer head 92,551 trainable parameters；相同 holdout 闭环 3/20。", "完成低秩视觉适配探索，但 3 cm / 8 of 20 部署门槛未通过；不部署。"],
            ["Frozen CLIP Patch Pointer", "393-scene pointer 版本在相同 20 个 holdout 中 2/20。", "增加 LoRA 后只有极小提升，未解决 patch-level 空间定位与接触执行。"],
            ["Contact monitor early regrasp", "seed-disjoint offline balanced accuracy 0.9196；闭环 V4 143/144 vs monitor 127/144；1 improvement, 17 regressions; p=0.000145。", "离线可分性不是安全的闭环干预条件；正式拒绝默认启用。"],
            ["Same-state early deep regrasp", "48 个 lift_post 分叉：continue V4 更优 47，early regrasp 更优 0，tie 1。", "不满足双向独有收益门槛；不继续训练 selector。"],
        ],
        [2150, 3650, 3560],
        compact=True,
    )
    add_text(document, "这组结果能够把论文的技术判断讲清：轻量化不等于可部署，离线准确率也不等于值得改变闭环轨迹。")

    document.add_heading("5.3 另外应保留的定量维度", level=2)
    for item in [
        "数据效率：10 / 25 / 50 / 92 demonstrations 下的 kNN、trajectory-kNN 与 object action head；只说 MuJoCo scripted demonstration 范围内的趋势。",
        "语言/空间：把闭集指令、颜色目标与“leftmost cube to white bowl”任务分别列出，不能把规则解析或闭集 CLIP 判别扩写为开放语义理解。",
        "资源：至少给 trainable parameters、训练时长、峰值 VRAM（存在记录时），让“轻量化”具有可审核的计算含义。",
        "扰动：MuJoCo friction、gain、force/gripper perturbation 作为 simulation robustness proxy，不能写成 Isaac 或真实机 transfer。",
    ]:
        add_bullet(document, item)

    document.add_heading("6. 论文图、表、视频的使用方案", level=1)
    add_table(
        document,
        ["论文材料", "建议位置", "要表达的唯一信息", "禁止误导"],
        [
            ["Figure 1: MuJoCo task scene", "Introduction / Environment", "对象、彩色目标盘、白碗、顶视相机和 WidowX 的任务边界。", "不要把仿真截图做成结果图。"],
            ["Figure 2: V4 execution pipeline", "Methods", "instruction -> frozen CLIP intent -> top RGB geometry -> structured executor -> one bounded retry。", "明确 state 仅用于 offline scoring。"],
            ["Table 1: task and metric protocol", "Methods", "任务、layout、seed、strict success、语言/对象选择指标。", "不要混合训练集与 holdout。"],
            ["Figure 3 / Table 2: result summary", "Results", "V4 两个 cohort、主要 baseline、LoRA 的闭环门槛结果。", "不要把视频矩阵的 27/27 illustrative clips 当 success rate。"],
            ["Figure 4: intervention / failure plot", "Discussion", "offline 0.9196 与闭环 regressions 的反差。", "不要只展示 contact monitor 的单条成功视频。"],
            ["Appendix evidence map", "Appendix", "方法 ID、评测文件、视频索引与复现命令。", "不要把整个文件夹截图当作证据。"],
        ],
        [1850, 1700, 3900, 1910],
        compact=True,
    )
    add_text(document, "视频展示的正确定位：你已有 3 methods x 3 tasks x 3 layouts = 27 条可溯源 MuJoCo 展示视频。它们适合说明同一个 task/layout 在不同方法下可切换观看；但这个矩阵被挑选为展示素材，不能用其中的成功数量替代实验总体成功率。")
    add_path(document, "showcase_assets/method_task_layout_matrix_v1/README.md; docs/video_evidence_index.md")
    document.add_heading("建议在论文正文或答辩中播放的 6 条视频", level=2)
    for item in [
        "V4：blue cube -> blue pad，standard layout，作为任务最直观的起点。",
        "V4：blue cube -> blue pad，repositioned 或 seven-object distractor layout，说明布局变化。",
        "V4：leftmost cube -> white bowl，说明关系型闭集指令任务。",
        "V4：一次 RGB table recovery 成功案例，配合首轮与最终 aggregate 成功率。",
        "CLIP LoRA Patch Pointer 成功案例，但视频标题必须同时标注 overall 3/20，避免 cherry-picking。",
        "Contact monitor false-trigger failure，配合 143/144 vs 127/144 的配对闭环统计，说明为何该策略被拒绝。",
    ]:
        add_bullet(document, item)

    document.add_heading("7. 海报应如何制作", level=1)
    add_text(document, "海报不是论文的缩印版。它的任务是让一位不了解项目的工程学院评审在约 90 秒内看懂：研究问题是什么、你测了什么、最终保留了什么、哪些路线被严谨地拒绝。尺寸、学院 logo 和提交格式必须以课程 brief 为准；若没有指定格式，可先使用 A0 横版三栏作为工作稿。")
    add_note(
        document,
        "推荐海报标题",
        "Lightweight Vision-Language-Action for MuJoCo WidowX Manipulation: Evidence from Adaptation Failures and RGB-Grounded Structured Control.",
        fill=PALE_GREEN,
    )
    add_table(
        document,
        ["版块", "推荐内容", "视觉形式"],
        [
            ["左栏：Problem and setup", "资源受限问题、WidowX 任务图、3 个代表任务、为何只使用 MuJoCo。", "一张场景图 + 3 行任务说明。"],
            ["中栏：Method and evidence", "V4 流程图；方法筛选线：BC -> temporal models -> PEFT/LoRA -> V4。", "一张管线图 + 小型证据时间线。"],
            ["右栏：Results and limits", "两个独立 cohort 主表；LoRA 3/20；contact monitor 被拒绝；结论和 scope。", "一个大数字、一个紧凑结果表、一个 failure comparison。"],
            ["底部：Reproducibility", "QR 指向 integrated research showcase 与 video matrix；方法、任务、布局切换入口。", "一个 QR + 1 句说明，避免自动播放六宫格。"],
        ],
        [1700, 4700, 2960],
        compact=True,
    )
    add_text(document, "海报上最值得放大的数字是“278/288 strict success”，但必须紧跟限定语：descriptive pooled replication across two seed-disjoint MuJoCo cohorts。这样展示结果强度，同时不掩盖统计边界。")
    add_text(document, "不要把所有 25 个方法、74 条视频或完整训练命令放在海报上；把它们放到 QR 所指向的展示页和附录索引。")

    document.add_heading("8. 答辩时的安全表述", level=1)
    add_table(
        document,
        ["被问到的问题", "建议的简洁回答"],
        [
            ["这是不是 VLA？", "它是资源受限的视觉-语言-动作研究与模块化 VLA baseline，而不是端到端 foundation VLA。语言、视觉定位和动作执行被明确分解。"],
            ["你是否完成了 LoRA？", "完成了 CLIP visual LoRA spatial pointer 的低秩适配探索，但 3/20 闭环成功未通过部署门槛。因此它是负结果，而不是最终策略。"],
            ["为什么 V4 成功率高？", "它没有让一个小模型直接回归厘米级抓取点，而是把闭集语义、RGB 几何、结构化接触执行和一次受限恢复分开处理；两组独立 seed 用相同协议复核。"],
            ["是否用了 MuJoCo 真值？", "运行时不读取物体真值来选择对象、决定重试或规划轨迹；真值只用于离线标签与严格评分。"],
            ["能迁移到真实机械臂吗？", "当前没有真实机器人验证，因此不作 transfer claim。项目保留的数据桥接、外部依赖审计与真实 WidowX 试验模板只属于后续工作。"],
            ["offline accuracy 0.9196 为什么不能用？", "它能预测标签，但将预测用于提前重抓后，独立闭环从 143/144 降至 127/144，并有 17 次回退。因此预测能力不足以证明干预价值。"],
        ],
        [2550, 6810],
        compact=True,
    )
    add_note(
        document,
        "一句话结论",
        "在本 MuJoCo 任务中，当前小数据和轻量适配不足以可靠地把 VLM 表征直接变成连续抓放动作；把语义、RGB 几何和接触执行分解，并以有界视觉恢复处理失败，形成了更可重复的工程基线。",
    )

    document.add_heading("9. 研究冻结后的四周写作计划", level=1)
    add_text(document, "这是写作和材料整理计划，不包含新实验。若 supervisor 给出不同截止日期，优先按其节奏调整。")
    add_table(
        document,
        ["周次", "交付物", "完成标准"],
        [
            ["Week 1", "论文目录、RQ、证据登记表、所有数字的单一来源链接。", "每条 claim 能指向一个 JSON/CSV/MD；删除无法证明的句子。"],
            ["Week 2", "Methods + Results 初稿；V4 主表、LoRA 负结果表、contact monitor 对照表。", "先有可审计主表，再选视频和截图。"],
            ["Week 3", "Introduction / Discussion / Conclusion；附录与引用检查。", "所有限制进入正文，不只放在最后一页。"],
            ["Week 4", "海报工作稿、90 秒版本、3-4 分钟版本、最终视频/QR 验证。", "每个视频与方法/任务/layout 对应，播放前确认统计脚注。"],
        ],
        [1050, 4520, 3790],
        compact=True,
    )

    document.add_heading("10. 提交前检查清单", level=1)
    for item in [
        "确认 programme handbook 要求：入学年度适用的 degree regulations、字数、论文模板、引用格式、提交系统和 poster brief。",
        "所有论文正文结果以原始 CSV/JSON 复核；把“成功率”限定到具体 cohort、任务和 strict metric。",
        "V4 结果写成 MuJoCo-only、seed-disjoint repeatability；不写成真实机、OpenVLA 或全局 SOTA 结论。",
        "LoRA、contact monitor 与 early regrasp 的 negative result 进入 Results/Discussion，而不是只留在附录。",
        "每个图表有编号、caption、数据来源和正文引用；每个视频都与方法、任务、layout 和结果记录一致。",
        "把完整方法门禁、运行命令、参数表、视频索引和文件路径放进 appendix 或 QR，而非堆在主文/海报。",
        "完成一次干净环境的复现检查：关键 viewer 命令、最终 JSON、展示网页与视频路径均可打开。",
    ]:
        add_bullet(document, item)

    document.add_heading("11. 推荐的本地证据入口", level=1)
    add_table(
        document,
        ["用途", "优先打开的文件"],
        [
            ["当前最终结论", "docs/final_closure_audit_v1.md; docs/final_research_closed_loop_v2.md"],
            ["独立 V4 复核", "docs/v4_independent_replication_v1.md"],
            ["LoRA negative result", "docs/clip_lora_patch_pointer_stage_v1.md"],
            ["接触监测器拒绝证据", "docs/contact_phase_monitor_heldout_v1_analysis.md"],
            ["每条 claim 的边界", "docs/claim_evidence_traceability.md"],
            ["方法与展示证据", "docs/method_evidence_gate.md; docs/video_evidence_index.md"],
            ["展示页面与视频矩阵", "docs/integrated_research_showcase.html; showcase_assets/method_task_layout_matrix_v1/README.md"],
            ["复现与附录", "docs/reproducible_command_index.md; docs/thesis_appendix_tables.md; docs/repository_reproducibility.md"],
        ],
        [2600, 6760],
        compact=True,
    )
    add_text(document, "University requirement note: 请以 University Regulations 的 current edition、入学年度适用的 degree regulations 和 programme handbook 确认所有正式要求；本文不替代课程或学院指示。")

    return document


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
