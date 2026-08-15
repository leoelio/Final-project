from __future__ import annotations

"""Build a thesis-writing guide from the completed MuJoCo research record."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "glasgow_msc_thesis_writing_guide_mujoco_vla.docx"

NAVY = "17365D"
BLUE = "2B6CB0"
PALE_BLUE = "EAF2F8"
PALE_GREY = "F3F6F8"
PALE_RED = "FDEDEC"
DARK = RGBColor(31, 41, 55)


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_margins(cell, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    properties.append(OxmlElement("w:cantSplit"))


def set_font(run, size: float, bold: bool = False, color: RGBColor | None = None, name: str = "Aptos") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_text(document: Document, text: str, style: str | None = None, bold: bool = False, color: RGBColor | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, 10.5, bold=bold, color=color or DARK)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(text)
    set_font(run, 10.2, color=DARK)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_font(run, 16 if level == 1 else 12.5, bold=True, color=RGBColor.from_string(NAVY if level == 1 else BLUE))


def add_note(document: Document, title: str, body: str, fill: str = PALE_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    prevent_row_split(table.rows[0])
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_font(run, 10.5, bold=True, color=RGBColor.from_string(NAVY))
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    run = paragraph.add_run(body)
    set_font(run, 10.0, color=DARK)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        if widths:
            cell.width = Inches(widths[index])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(header)
        set_font(run, 9.2, bold=True, color=RGBColor(255, 255, 255))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_index, row in enumerate(rows):
        table_row = table.add_row()
        prevent_row_split(table_row)
        cells = table_row.cells
        for index, value in enumerate(row):
            cell = cells[index]
            if widths:
                cell.width = Inches(widths[index])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, "FFFFFF" if row_index % 2 == 0 else PALE_GREY)
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            run = paragraph.add_run(value)
            set_font(run, 8.8, color=DARK)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_pipeline(document: Document, labels: list[str], caption: str) -> None:
    table = document.add_table(rows=1, cols=len(labels))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    label_width = 6.5 / len(labels)
    for index, label in enumerate(labels):
        cell = table.cell(0, index)
        cell.width = Inches(label_width)
        set_cell_shading(cell, PALE_BLUE)
        set_cell_margins(cell, top=180, bottom=180, start=80, end=80)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        set_font(run, 8.3, bold=True, color=RGBColor.from_string(NAVY))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run(caption)
    set_font(run, 9, color=RGBColor(90, 100, 110))


def add_template(document: Document, english: str, guidance: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    prevent_row_split(table.rows[0])
    cells = table.rows[0].cells
    for cell, fill in zip(cells, (PALE_GREY, "FFFFFF")):
        set_cell_shading(cell, fill)
        set_cell_margins(cell, top=110, bottom=110, start=140, end=140)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[0].width = Inches(3.0)
    cells[1].width = Inches(3.7)
    run = cells[0].paragraphs[0].add_run(english)
    set_font(run, 9.2, bold=True, color=RGBColor.from_string(NAVY))
    run = cells[1].paragraphs[0].add_run(guidance)
    set_font(run, 9.2, color=DARK)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def page_break(document: Document) -> None:
    document.add_section(WD_SECTION.NEW_PAGE)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("MuJoCo WidowX MSc Thesis Writing Guide | Working document | 30 July 2026")
    set_font(run, 8, color=RGBColor(100, 110, 120))


def cover(document: Document) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(30)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("UNIVERSITY OF GLASGOW\nCOLLEGE OF SCIENCE AND ENGINEERING")
    set_font(run, 13, bold=True, color=RGBColor.from_string(NAVY))
    paragraph.paragraph_format.space_after = Pt(32)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("MSc Dissertation Writing Guide")
    set_font(run, 24, bold=True, color=RGBColor.from_string(NAVY))
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Lightweight Vision-Language-Action Methods for MuJoCo WidowX Desktop Manipulation")
    set_font(run, 15, bold=True, color=RGBColor.from_string(BLUE))
    paragraph.paragraph_format.space_after = Pt(28)
    add_note(
        document,
        "Purpose",
        "A project-specific writing blueprint. It transforms the recorded MuJoCo experiments into an evidence-led engineering dissertation. It is not an official University template and must be checked against the current programme handbook, project brief, supervisor instructions and submission portal.",
    )
    add_table(
        document,
        ["Field", "Recommended entry"],
        [
            ["Candidate", "[Your full name and student number]"],
            ["Programme", "[Exact MSc programme title]"],
            ["Supervisor", "[Supervisor name]"],
            ["Submission", "[Month Year]"],
            ["Working title", "Resource-Constrained Vision-Language-Action for MuJoCo WidowX Desktop Manipulation"],
        ],
        [1.5, 5.2],
    )
    add_text(document, "Prepared from the completed MuJoCo research package. All numerical claims in this guide should be cross-checked against the linked experiment records before submission.", bold=True, color=RGBColor.from_string(NAVY))


def build_document() -> Document:
    document = Document()
    configure_document(document)
    cover(document)
    page_break(document)

    add_heading(document, "1. How to Use This Guide")
    add_text(document, "This document is a chapter-by-chapter writing plan, not a dissertation to submit unchanged. Draft the thesis in English, use your programme's approved template, and replace all bracketed placeholders with verified material. Each claim must point to a recorded CSV, JSON, figure, table or video index in the project workspace.")
    add_note(document, "University check", "The University of Glasgow 2025-26 College of Science and Engineering MSc regulations describe a 180-credit award including a supervised independent work component of 60 credits or more. The exact dissertation format, word limit, front matter, ethics declaration, deadline and marking criteria come from your programme handbook and project course documentation. Check these before writing the final version.")
    add_table(document, ["Writing rule", "Application to this project"], [
        ["Separate evidence from aspiration", "Write completed MuJoCo experiments in past tense. Write OpenVLA LoRA, Isaac and real WidowX work only as future work."],
        ["Make the scope explicit", "The formal deliverable is MuJoCo-only, not sim-to-real or a real-robot study."],
        ["Use aggregate results", "Report fixed-seed tables and confidence intervals; do not use individual videos as proof of performance."],
        ["Report negative results", "The contact-monitor and early-regrasp regressions are evidence about deployment safety, not material to hide."],
    ], [1.7, 5.0])

    add_heading(document, "2. Thesis Positioning and Proposed Title")
    add_text(document, "Recommended title: Resource-Constrained Vision-Language-Action for MuJoCo WidowX Desktop Manipulation: A Study of Lightweight Adaptation, RGB Grounding and Structured Control.")
    add_text(document, "Recommended central argument: under the available data and compute constraints, direct continuous-action imitation and lightweight VLM/LoRA-style spatial heads did not provide robust closed-loop grasping. A decomposed controller that assigns intent to frozen CLIP, geometry to top-view RGB, and contact execution to a structured controller produced a reproducible MuJoCo baseline.")
    add_template(document, "This dissertation investigates whether lightweight post-training can improve language-conditioned robot manipulation under constrained compute and demonstration budgets.", "Opening sentence for the abstract or introduction. Follow with the MuJoCo-only scope immediately.")
    add_template(document, "The contribution is an evidence-led comparison rather than a claim of a fully trained robot foundation model.", "Use this to control examiner expectations and distinguish the project from OpenVLA-scale studies.")

    add_heading(document, "3. Recommended Dissertation Structure")
    add_table(document, ["Chapter", "Suggested allocation", "Purpose", "Project evidence"], [
        ["Abstract", "250-350 words", "Question, method, result, boundary", "V4 278/288; no real OpenVLA claim"],
        ["1 Introduction", "10-12%", "Motivation, gap, RQs, contributions", "Resource constraint and task design"],
        ["2 Related Work", "15-18%", "VLA, IL, ACT, diffusion, PEFT, sim evaluation", "Literature, not local reports"],
        ["3 Problem Formulation", "8-10%", "Define observation, action, success and constraints", "Task and metrics definitions"],
        ["4 Methodology", "18-22%", "Environment, methods, data, evaluation protocol", "Environment/code/pipeline"],
        ["5 Results and Analysis", "22-26%", "Main results, ablations, negative results", "CSV/JSON/video indexes"],
        ["6 Conclusion and Future Work", "7-9%", "Answer RQs and delimit scope", "Final closure audit"],
        ["Appendices", "not counted unless handbook says otherwise", "Commands, long tables, prompts, protocols", "Method matrix and reproducibility records"],
    ], [1.3, 1.15, 2.2, 2.1])
    add_note(document, "Word-count caution", "The suggested proportions are writing advice only. Do not infer a University of Glasgow MSc word limit from this guide; use the current programme documentation and supervisor direction.", PALE_RED)

    add_heading(document, "4. System Architecture")
    add_pipeline(document, ["1. Language\nprompt", "2. Frozen CLIP\nintent", "3. Top RGB\ngeometry", "4. Pick/place\ncontrol", "5. One RGB\nretry", "6. Offline\nscore"], "Figure 4.1. Final MuJoCo V4 execution architecture. Read left to right; draw this as a clean vector figure in the dissertation.")
    add_text(document, "Write one paragraph below the architecture that assigns responsibility precisely: CLIP performs four-class closed-vocabulary intent recognition; top-view RGB estimates object/target geometry; the structured controller produces the contact trajectory; MuJoCo state is reserved for offline evaluation. This protects the thesis from an invalid claim that simulation ground truth was used for deployment decisions.")
    add_table(document, ["Module", "Input", "Output", "Claim boundary"], [
        ["Intent module", "Instruction and initial RGB", "Task intent", "Frozen CLIP representation with a light semantic adapter; not a robot VLA"],
        ["Grounding module", "Top RGB + calibration", "Source/target table coordinates", "Image geometry, not privileged MuJoCo state at run time"],
        ["Execution module", "Coordinates + task configuration", "7-D joint-position controls", "Structured expert, not learned continuous control"],
        ["Recovery module", "Post-attempt RGB", "At most one re-localised retry", "Bounded rule; not a learned recovery policy"],
        ["Evaluator", "MuJoCo trajectory/state", "Strict success, distance, grasp audit", "Offline scoring only"],
    ], [1.35, 1.35, 1.4, 2.65])

    add_heading(document, "5. Workflow and Experimental Process")
    add_pipeline(document, ["1. Define\ntask", "2. Collect\ndemos", "3. Replay\naudit", "4. Train\nbaseline", "5. Fixed-seed\nevaluation", "6. Aggregate\nreport", "7. Viewer /\nvideo"], "Figure 4.2. Research workflow. Read left to right and keep evaluation seeds frozen before comparing candidates.")
    add_text(document, "Use this workflow in Methodology, not as a narrative diary. Explain that each method version retains a training record, an evaluation protocol, a viewer command and a video index. Then explain that candidate changes were accepted or rejected using closed-loop results rather than visual impressions alone.")

    add_heading(document, "6. Introduction: What to Write")
    add_text(document, "The introduction should move from the practical problem to the precise research question in four pages or fewer. Avoid a marketing-style introduction. Start with the cost of robot data and the mismatch between foundation-model scale and a student project’s compute/data budget.")
    add_table(document, ["Paragraph", "Required content", "English prompt"], [
        ["P1: motivation", "Language-conditioned manipulation needs semantic understanding and contact-stable control.", "Robotic manipulation systems must combine task interpretation with reliable physical execution."],
        ["P2: gap", "Large VLA reports do not directly answer what works with limited data/compute.", "The practical question is not whether large VLA models can scale, but which components remain useful under constrained resources."],
        ["P3: research question", "State the MuJoCo WidowX setting, tasks and comparison.", "Can lightweight visual-language adaptation outperform ordinary imitation-learning baselines in a controlled MuJoCo tabletop setting?"],
        ["P4: contributions", "List environment, comparison, final modular baseline and negative results.", "This work makes four contributions: ..."],
        ["P5: roadmap", "End with chapter map.", "The remainder of this dissertation is organised as follows..."],
    ], [1.2, 2.55, 2.95])
    add_text(document, "Suggested contributions: (1) a reproducible MuJoCo WidowX benchmark with semantic, colour-target and spatial tasks; (2) a controlled comparison of conventional IL and lightweight VLA-style proxies; (3) a run-time RGB-grounded, structured controller with independent V4 replication; and (4) negative closed-loop evidence showing why selected spatial and contact-learning candidates were rejected.")

    add_heading(document, "7. Related Work: Literature Map")
    add_text(document, "Use peer-reviewed papers and original project repositories as the primary sources. Do not cite the project’s own experiment reports as related work. The section should compare mechanisms, assumptions and compute/data requirements rather than list model names.")
    add_table(document, ["Subsection", "Questions to answer", "How this project differs"], [
        ["Vision-language-action", "How do VLA models map image/language histories to robot actions?", "This work uses VLA-inspired lightweight proxies and a modular final system, not a full pretrained robot VLA."],
        ["Imitation learning", "Why do BC, action chunks and retrieval policies suffer from distribution shift?", "Closed-loop failures of linear/MLP/ACT-lite serve as local diagnostic evidence."],
        ["ACT and diffusion", "What do temporal modelling and action chunks solve?", "State-only/lite baselines are included; do not claim official ACT or visual Diffusion Policy reproduction."],
        ["PEFT and LoRA", "Why freeze large encoders and adapt small modules?", "The project tests LoRA-style and CLIP-LoRA proxies; it does not complete OpenVLA LoRA."],
        ["Simulated manipulation", "What are the roles and limitations of MuJoCo, domain randomisation and sim-to-real?", "The reported setting stops at MuJoCo; no physical transfer claim is made."],
    ], [1.4, 2.8, 2.5])
    add_note(document, "Citation discipline", "Every factual claim in Related Work needs a primary citation. Describe model sizes, training data and hardware only when the cited source supports them. Keep a BibTeX database and cite the official OpenVLA, ACT, Diffusion Policy, CLIP and MuJoCo sources directly.")

    add_heading(document, "8. Problem Formulation and Mathematical Language")
    add_text(document, "Use notation only when it clarifies a design choice. Define every symbol before use and connect each equation to an implemented module. The following formulation is sufficient for this project.")
    add_table(document, ["Concept", "Suggested notation", "Interpretation"], [
        ["Instruction", "l", "Natural-language task instruction"],
        ["RGB observation", "I_t", "Top-view image at simulation step t"],
        ["Robot state", "q_t, qdot_t, g_t", "Joint position, velocity and gripper state"],
        ["Observation", "o_t = (I_t, q_t, qdot_t, g_t, l)", "Available sensory/task input"],
        ["Action", "a_t in R^7", "Joint-position targets and gripper control"],
        ["Policy", "a_t = pi_theta(o_{t-h:t})", "Learned policy for baseline methods"],
        ["Task success", "S(tau) in {0,1}", "Strict lift, placement and boundary criteria"],
    ], [1.45, 2.05, 3.2])
    add_text(document, "For behaviour cloning, present: L_BC(theta) = (1/N) sum_i ||a_i - pi_theta(o_i)||^2. Then explain why low offline action MSE did not guarantee closed-loop grasping, using linear BC as evidence. For action chunks, define A_t = [a_t, ..., a_(t+H-1)] and explain temporal execution/replanning.")
    add_text(document, "For the semantic adapter, define p_phi(y | I_0, l) over the four task intents. For LoRA-style discussion, use W' = W + BA with rank r much smaller than min(d_in, d_out). State that this is the parameter-efficient adaptation principle; the completed project contains proxy experiments and not a full OpenVLA LoRA result.")
    add_text(document, "For the final controller, define a composed policy pi(o_t, l) = E(G(I_t, p_phi(l, I_0))), where G returns geometric source/target coordinates and E is the structured executor. Define the retry rule as a bounded transition that may occur at most once after a failed first attempt. Do not present this as an end-to-end differentiable policy.")

    add_heading(document, "9. Methodology")
    add_heading(document, "9.1 Environment, Tasks and Data", level=2)
    add_text(document, "Describe the WX250S model, tabletop MJCF scene, top and front RGB cameras, seven controlled action dimensions, object set, target pads and bowl. Use a task table and a scene figure. State that scripted expert rollouts generated demonstrations and each stored seed, object poses, action trajectory, success/attempt labels and replayable `.npz` files.")
    add_table(document, ["Task", "Semantic challenge", "Control challenge", "Use in thesis"], [
        ["Blue cube -> blue pad", "Basic colour grounding", "Pick and place", "Core placement task"],
        ["Blue cube -> red pad", "Source/target colour distinction", "Transfer and release", "Target-region generalisation"],
        ["Red cube -> red pad", "Same-colour source/target disambiguation", "Object identity", "Visual exclusion analysis"],
        ["Leftmost cube -> bowl", "Spatial relation", "Selection under distractors", "Language/spatial task"],
    ], [1.6, 2.2, 1.65, 1.25])
    add_heading(document, "9.2 Methods and Baselines", level=2)
    add_text(document, "Group methods by purpose instead of forcing all 25 versions into the main text. Put the full 25-method table in an appendix. In the main chapter show: scripted expert/structured controller; ordinary BC family; trajectory/action-chunk family; lightweight PEFT/VLM proxies; direct CLIP spatial heads; and the final RGB-grounded controller.")
    add_heading(document, "9.3 Evaluation Protocol", level=2)
    add_text(document, "Define train-range, held-out, language/OOD and contact-domain protocols. Use task success, strict grasp success, semantic correctness, initial object selection, target distance, trainable parameters, training time and peak VRAM where recorded. State seed ranges before results, not after observing outcomes.")

    add_heading(document, "10. Results, Analysis and Ablation")
    add_heading(document, "10.1 Main Result Narrative", level=2)
    add_text(document, "Lead the results chapter with the question, the protocol and the aggregate table. The final V4 system achieved 135/144 (93.8%) in one independent batch and 143/144 (99.3%) in a second disjoint batch; the descriptive combined result is 278/288 (96.5%), Wilson 95% CI 0.937-0.981. Semantic correctness and initial visual object selection were 288/288 within this defined evaluation scope.")
    add_note(document, "Interpretation", "This result supports repeatability of the final modular MuJoCo controller within the recorded tasks and contact domains. It does not establish a causal superiority claim over all learned methods, and it must not be reported as a real-robot or end-to-end VLA result.")
    add_heading(document, "10.2 Baseline Comparison", level=2)
    add_table(document, ["Family", "Representative evidence", "What to conclude"], [
        ["Ordinary BC", "Linear BC 0/5 held-out; MLP BC 1/5 held-out", "Small-data one-step action regression is not closed-loop robust."],
        ["Memory/retrieval", "kNN BC 5/5 train-range, 1/5 held-out; trajectory kNN 5/5 train-range, 0/5 held-out", "Local demonstration retrieval can memorise but did not generalise."],
        ["ACT/diffusion lite", "State-only/lite versions did not form stable held-out grasping", "More temporal modelling alone did not solve contact instability."],
        ["Lightweight action heads", "Object-language, Adapter and LoRA-style proxies: 0-1/5 held-out", "Parameter efficiency alone did not solve centimetre-scale spatial control."],
        ["Direct CLIP spatial heads", "Frozen patch pointer 3/20; CLIP-LoRA patch pointer 3/20", "Better representation/point error did not pass the closed-loop deployment gate."],
    ], [1.55, 2.6, 2.55])
    add_heading(document, "10.3 Ablation Plan", level=2)
    add_text(document, "Write ablations as controlled questions. Use an explicit reference method and change one factor per row. The table below is a valid dissertation structure; populate it only from completed records.")
    add_table(document, ["Ablation question", "Controlled comparison", "Reported evidence", "Correct conclusion"], [
        ["Does RGB retry help?", "V4 first attempt vs one bounded retry", "127/144 first-round vs 135/144 after retry in V4 batch 1; 8 recoveries, no regressions", "Bounded retry shows positive within-protocol recovery evidence."],
        ["Does wider table search help?", "Source search vs table search, same seeds/actions", "133/144 vs 135/144; p=0.5000", "Positive candidate only; not statistically significant."],
        ["Does CLIP LoRA solve spatial grounding?", "Frozen and LoRA patch-pointer heads", "Both 3/20 closed-loop in the recorded gate", "Do not promote to the deployed controller."],
        ["Does early contact regrasp help?", "V4 standard vs monitor-triggered early regrasp", "143/144 vs 127/144; 1 improvement, 17 regressions; p=0.000145", "Reject the trigger for default deployment."],
    ], [1.65, 2.0, 2.0, 1.05])
    add_heading(document, "10.4 Failure Analysis", level=2)
    add_text(document, "Classify failures by perception, semantic intent, grasp/contact, transport and terminal verification. The key analytic insight is that high offline monitor accuracy (0.9196 balanced accuracy) did not translate into a safe intervention policy. Explain the distinction between a predictive label and the counterfactual value of changing a successful trajectory.")
    add_text(document, "Include one success video frame and one failure frame only as qualitative illustrations. Reference the aggregate statistical table for every numerical performance claim.")

    add_heading(document, "11. Conclusion")
    add_template(document, "This dissertation evaluated lightweight visual-language adaptation for MuJoCo WidowX manipulation under constrained data and compute.", "First sentence of the conclusion.")
    add_template(document, "The evidence indicates that decomposing semantic intent, visual geometry and contact execution was more reliable than directly regressing continuous actions from the available lightweight baselines.", "Answer the main research question without overstating generality.")
    add_template(document, "The main limitation is that the evidence is confined to a MuJoCo environment and does not demonstrate real-robot transfer or full OpenVLA fine-tuning.", "Mandatory scope sentence.")
    add_text(document, "End with three concise contributions, the answer to each research question, and one sentence on reproducibility: code, scene assets, evaluation scripts, small runtime assets and a private source repository are organised for rerunning the core system.")

    add_heading(document, "12. Future Work")
    add_table(document, ["Priority", "Next experiment", "Required evidence before claiming success"], [
        ["1", "Collect more diverse contact trajectories, including counterfactual-safe interventions.", "New fixed-seed closed-loop comparisons with regressions as a primary rejection metric."],
        ["2", "Use a compatible 27GB+ GPU for a genuine VLA/LoRA experiment.", "RLDS/action-normalised dataset, exact base model/version, GPU/VRAM report and held-out videos."],
        ["3", "Test a small visual ACT or SmolVLA-style model with a compute budget study.", "Same task protocol, same metrics and model/resource table."],
        ["4", "Introduce real WidowX validation only after simulation protocols are frozen.", "New real-robot safety protocol, calibration, data and separately reported results."],
        ["5", "Extend language coverage beyond four closed intents.", "Pre-registered paraphrase/compositional OOD protocol and semantic error analysis."],
    ], [0.8, 3.0, 2.9])

    add_heading(document, "13. Appendices and Evidence Pack")
    add_text(document, "Use appendices to make the dissertation auditable without interrupting the main argument. The appendix should be referred to in the main text, not treated as an unstructured file dump.")
    add_table(document, ["Appendix", "Contents", "Existing source"], [
        ["A: Environment", "MJCF scene, action space, task specification and object layout", "assets/, widowx_env/tabletop_env.py"],
        ["B: Methods", "Full 25-method table, parameter counts and boundary notes", "docs/thesis_appendix_tables.md"],
        ["C: Protocol", "Seeds, success definition, strict grasp audit and command index", "docs/reproducible_command_index.md"],
        ["D: Extra results", "Per-task/OOD tables, resource table and domain tests", "docs/model_resource_summary.md and CSVs"],
        ["E: Reproducibility", "Repository structure, requirements and final viewer command", "README.md and docs/repository_reproducibility.md"],
        ["F: Video evidence", "Short curated playlist and captions", "docs/defense_video_playlist.html"],
    ], [1.15, 3.05, 2.5])

    add_heading(document, "14. Figure and Table Checklist")
    add_table(document, ["Item", "Where", "Purpose"], [
        ["Figure 1", "Introduction", "MuJoCo tabletop scene with labelled objects, targets and camera."],
        ["Figure 2", "Methodology", "Final system architecture: language -> CLIP -> RGB -> expert -> retry."],
        ["Figure 3", "Methodology", "Research workflow from demonstration collection to evaluation."],
        ["Table 1", "Methodology", "Task matrix and success definition."],
        ["Table 2", "Results", "Grouped baseline results: train/held-out/language/resource."],
        ["Table 3", "Results", "V4 independent replication and confidence interval."],
        ["Table 4", "Analysis", "Ablation and rejected-candidate interventions."],
        ["Figure 4", "Analysis", "Failure taxonomy or paired improvement/regression plot."],
    ], [1.1, 1.45, 4.15])

    add_heading(document, "15. Final Submission Checklist")
    for item in [
        "Use the current programme handbook as the final authority for word count, formatting, submission route and marking criteria.",
        "Check every number against the source CSV/JSON; never copy a success rate from a single video caption.",
        "Use past tense for completed experiments and future tense for OpenVLA, Isaac and real-robot plans.",
        "Define the strict success metric before presenting performance results.",
        "Use a reference manager and verify every in-text citation has a bibliography entry.",
        "Check all figures have captions, readable labels and are cited in the text before they appear.",
        "Place full commands, model paths and long method tables in appendices; keep the main narrative analytical.",
        "Run the repository bundle verification and final viewer command before the final defence/demo.",
    ]:
        add_bullet(document, item)

    add_heading(document, "16. Source Records Used by This Guide")
    add_text(document, "Project evidence is indexed in docs/, especially final_closure_audit_v1.md, final_research_closed_loop_v2.md, thesis_appendix_tables.md and repository_reproducibility.md. Verify University requirements against the current programme handbook and https://www.gla.ac.uk/myglasgow/apg/policies/uniregs/regulations2025-26/scieng/msceng/.")
    return document


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
