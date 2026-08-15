from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "mujoco_html_3to4min_presentation_guide_zh.docx"

BLUE = "2E74B5"
DARK = "0B2545"
TEAL = "287F69"
MUTED = "5B6573"
WHITE = "FFFFFF"
LIGHT_BLUE = "E8EEF5"
LIGHT_GREEN = "E7F3EE"
LIGHT_GOLD = "FFF5D6"
LIGHT_RED = "FBEAEC"


def font(run, size=11, color=DARK, bold=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold


def insert_before(parent, element, following_tags):
    for tag in following_tags:
        following = parent.find(qn(tag))
        if following is not None:
            following.addprevious(element)
            return
    parent.append(element)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        insert_before(tc_pr, node, ("w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign"))
    node.set(qn("w:fill"), fill)


def cell_format(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        insert_before(tc_pr, borders, ("w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign"))
    for edge in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:color"), "D4DCE7")
    mar = tc_pr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        insert_before(tc_pr, mar, ("w:textDirection", "w:tcFitText", "w:vAlign"))
    for side, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        value_node = mar.find(qn(f"w:{side}"))
        if value_node is None:
            value_node = OxmlElement(f"w:{side}")
            mar.append(value_node)
        value_node.set(qn("w:w"), str(value))
        value_node.set(qn("w:type"), "dxa")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    insert_before(table_pr, tbl_ind, ("w:tblBorders", "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"))
    for grid, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid.set(qn("w:w"), str(width))
    for row in table.rows:
        row_pr = row._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell_format(cell)


def p(doc, text="", size=11, color=DARK, bold=False, align=None, before=0, after=6, style=None, keep=False):
    para = doc.add_paragraph(style=style)
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.keep_with_next = keep
    if align is not None:
        para.alignment = align
    if text:
        font(para.add_run(text), size, color, bold)
    return para


def heading(doc, text, level=1):
    return p(doc, text, style=f"Heading {level}", keep=True)


def callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    font(para.add_run(f"{label}  "), 10.5, DARK, True)
    font(para.add_run(text), 10.5, DARK)
    p(doc, "", after=2)


def speech(doc, time, text):
    table = doc.add_table(rows=1, cols=2)
    set_table(table, [1250, 8110])
    left, right = table.rows[0].cells
    shade(left, TEAL)
    left_p = left.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_p.paragraph_format.space_after = Pt(0)
    font(left_p.add_run(time), 10, WHITE, True)
    shade(right, LIGHT_BLUE)
    right_p = right.paragraphs[0]
    right_p.paragraph_format.space_after = Pt(0)
    font(right_p.add_run("可直接讲："), 10.5, DARK, True)
    font(right_p.add_run(f"“{text}”"), 10.5, DARK)
    p(doc, "", after=2)


def time_table(doc):
    table = doc.add_table(rows=1, cols=3)
    set_table(table, [1350, 2560, 5450])
    headers = ("时间", "网页位置", "只讲这一句话")
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, BLUE)
        para = cell.paragraphs[0]
        para.paragraph_format.space_after = Pt(0)
        font(para.add_run(text), 9.5, WHITE, True)
    rows = (
        ("0:00-0:30", "研究问题", "先定义范围：MuJoCo 的轻量语义适配，不是端到端 VLA。"),
        ("0:30-1:20", "标准搬运对比", "0/20 与 20/20 诊断的是接触控制分工。"),
        ("1:20-2:15", "低摩擦压力测试", "30/40 到 36/40 是渐变趋势，不是显著结论。"),
        ("2:15-2:50", "语言 + 资源", "闭词表内新句式 40/40，4,100 参数 adapter。"),
        ("2:50-3:40", "结论 / 可选视频", "一句话收束；有余量才放蓝方块到蓝盘。"),
    )
    for values in rows:
        cells = table.add_row().cells
        row_pr = table.rows[-1]._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        for index, (cell, text) in enumerate(zip(cells, values)):
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            if index == 0:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            font(para.add_run(text), 9.2, DARK)
    p(doc, "", after=2)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after in (("Heading 1", 16, 18, 10), ("Heading 2", 13, 14, 7)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(header.add_run("MuJoCo-only | 3-4 分钟答辩稿"), 8.5, MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("毕业设计演讲准备"), 8.5, MUTED)


def build():
    doc = Document()
    configure(doc)
    p(doc, "MuJoCo 轻量视觉-语言机械臂操作", 22, DARK, True, WD_ALIGN_PARAGRAPH.CENTER, before=24, after=4)
    p(doc, "3-4 分钟 HTML 演讲稿", 15, BLUE, True, WD_ALIGN_PARAGRAPH.CENTER, after=12)
    callout(doc, "使用规则", "按下面的时间段讲；视频是可选项。全程只说 MuJoCo-only、冻结 CLIP 意图适配和结构化执行，不说端到端 VLA、LoRA 或真实机械臂验证。", LIGHT_GOLD)
    heading(doc, "一页时间线", 1)
    time_table(doc)
    heading(doc, "可直接照着讲", 1)
    speech(doc, "0:00-0:30", "我的研究问题是在有限示范和算力下，视觉语言表征能否帮助 WidowX 完成桌面搬运。当前答案是可以完成任务级语义选择，但可靠抓取来自结构化接触执行。因此这是 MuJoCo 的轻量语义适配原型，不是端到端 VLA。")
    speech(doc, "0:30-1:20", "请看标准搬运对比。严格成功要求语义正确、抬升至少 0.06 米并持续接近 TCP，最后放入目标区。冻结 CLIP 连续动作头是 0/20，采用的分层方案是 20/20。这个差异说明图文表征不等于接触控制；把语义选择和抓取执行解耦后，才形成可靠搬运。")
    speech(doc, "1:20-2:15", "0 和 100 是严格门槛下的饱和结果，所以我继续在低摩擦压力下做配对测试。四个任务共 40 条配对轨迹，标准执行器从 30/40，融合执行器到 36/40，平均最终目标距离从 0.0568 米降到 0.0252 米。最左方块到碗没有成功率提升，因此这只是物理鲁棒性的改善趋势。精确 McNemar 检验 p 等于 0.0703，不能称为统计显著。")
    speech(doc, "2:15-2:50", "语言方面，采用闭词表规范化后，固定词表内的 8 条新句子达到 40/40。它证明的是词表内新句式，不是开放词汇泛化。资源方面，151M 参数的 CLIP 编码器保持冻结，只训练 4,100 个意图 adapter 参数，用 79 条成功示范完成四类任务选择。")
    speech(doc, "2:50-3:30", "最后总结：本项目证明了轻量视觉语言模块可以在小规模 MuJoCo 桌面任务中负责语义选择，但接触可靠性仍需要结构化执行和反馈。下一步是把抓取和抬升改成可学习的闭环子策略，并扩大独立 OOD 和扰动评测。")
    heading(doc, "可选的 8 秒视频", 1)
    callout(doc, "只在时间充足时播放", "播放“蓝方块 -> 蓝盘”视频。播放前说“下面只用一条视频核对数据对应的实际过程”，物体落入蓝盘后马上回到结论。视频不是成功率来源，成功率来自批量评测。", LIGHT_GREEN)
    heading(doc, "最后一定要守住的三条边界", 1)
    callout(doc, "1. 方法边界", "不是端到端 VLA、LoRA 或 RL 后训练；抓取控制由结构化 waypoint 执行。", LIGHT_RED)
    callout(doc, "2. 统计边界", "30/40 -> 36/40 是低摩擦下的改善趋势；p=0.0703，不显著。", LIGHT_RED)
    callout(doc, "3. 外推边界", "所有结论只限 MuJoCo 的四类任务和固定闭词表；没有真实机械臂结果。", LIGHT_RED)
    p(doc, "对应网页：docs\\mujoco_research_summary_zh_en.html", 9, MUTED, False, before=8, after=0)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
