from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "mujoco_html_presentation_guide_zh.docx"

BLUE = "2E74B5"
DARK_BLUE = "0B2545"
TEAL = "287F69"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_GREEN = "E7F3EE"
PALE_GOLD = "FFF5D6"
PALE_RED = "FBEAEC"
MUTED = "5B6573"
WHITE = "FFFFFF"


def set_run_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def insert_before(parent, element, following_tags):
    for tag in following_tags:
        following = parent.find(qn(tag))
        if following is not None:
            following.addprevious(element)
            return
    parent.append(element)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        insert_before(tc_pr, shd, ("w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign"))
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        insert_before(tc_pr, tc_mar, ("w:textDirection", "w:tcFitText", "w:vAlign"))
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="D4DCE7", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        insert_before(tc_pr, borders, ("w:shd", "w:noWrap", "w:tcMar", "w:textDirection", "w:tcFitText", "w:vAlign"))
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        insert_before(table_pr, tbl_ind, ("w:tblBorders", "w:shd", "w:tblLayout", "w:tblCellMar", "w:tblLook"))
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths_dxa):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph):
    paragraph.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    paragraph.add_run(" 页")


def add_paragraph(doc, text="", style=None, size=None, color=None, bold=None, italic=None,
                  align=None, before=None, after=None, keep_with_next=False):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    if before is not None:
        paragraph.paragraph_format.space_before = Pt(before)
    if after is not None:
        paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.keep_with_next = keep_with_next
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_heading(doc, text, level=1):
    style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}[level]
    return add_paragraph(doc, text, style=style, keep_with_next=True)


def add_label_paragraph(doc, label, text, fill=None):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    if fill:
        set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}  ")
    set_run_font(label_run, size=10.5, color=DARK_BLUE, bold=True)
    value_run = p.add_run(text)
    set_run_font(value_run, size=10.5, color=DARK_BLUE)
    add_paragraph(doc, "", after=2)
    return table


def add_script(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run("可直接讲：")
    set_run_font(lead, size=10.5, color=DARK_BLUE, bold=True)
    quote = p.add_run(f"“{text}”")
    set_run_font(quote, size=10.5, color=DARK_BLUE)
    add_paragraph(doc, "", after=2)
    return table


def write_cell(cell, text, header=False, color=None, bold=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=9.5):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, color=color or (WHITE if header else DARK_BLUE), bold=header if bold is None else bold)
    if header:
        set_cell_shading(cell, BLUE)


def add_data_table(doc, headers, rows, widths_dxa, header_fill=BLUE, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    prevent_row_split(hdr)
    for cell, header in zip(hdr.cells, headers):
        write_cell(cell, header, header=True, size=font_size)
        set_cell_shading(cell, header_fill)
    for row_values in rows:
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for index, (cell, value) in enumerate(zip(cells, row_values)):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index > 0 else WD_ALIGN_PARAGRAPH.LEFT
            write_cell(cell, str(value), size=font_size, align=alignment)
    add_paragraph(doc, "", after=2)
    return table


def add_question(doc, question, answer):
    p = add_paragraph(doc, before=5, after=2, keep_with_next=True)
    label = p.add_run(f"问：{question}")
    set_run_font(label, size=11, color=DARK_BLUE, bold=True)
    p = add_paragraph(doc, after=6)
    label = p.add_run("答：")
    set_run_font(label, size=11, color=TEAL, bold=True)
    answer_run = p.add_run(answer)
    set_run_font(answer_run, size=11, color=DARK_BLUE)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    header_run = header_p.add_run("MuJoCo-only WidowX 桌面操作 | HTML 答辩演讲指南")
    set_run_font(header_run, size=8.5, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    footer_run = footer_p.add_run("毕业设计演讲准备 | ")
    set_run_font(footer_run, size=8.5, color=MUTED)
    add_page_number(footer_p)
    for run in footer_p.runs:
        set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc):
    add_paragraph(doc, "毕业设计答辩辅助材料", size=11, color=TEAL, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, before=70, after=12)
    add_paragraph(doc, "MuJoCo 轻量视觉-语言机械臂操作", size=24, color=DARK_BLUE, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=7)
    add_paragraph(doc, "HTML 研究结果页讲解与答辩指南", size=18, color=BLUE, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=24)
    add_paragraph(doc, "对应页面：mujoco_research_summary_zh_en.html", size=10.5, color=MUTED,
                  align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    add_paragraph(doc, "用途：将网页中的指标、视频和边界转化为一条可直接讲述的研究叙事。",
                  size=11, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
    add_label_paragraph(
        doc,
        "核心定位",
        "这是一项 MuJoCo-only 的 WidowX 桌面搬运研究：冻结 CLIP 负责四类任务意图，结构化 waypoint 与接触反馈负责抓取、抬升、放置和有限恢复。",
        LIGHT_BLUE,
    )
    add_label_paragraph(
        doc,
        "建议时长",
        "8 至 10 分钟讲解：数据与结论为主，视频只作短证据；把每条定量结论限制在网页给出的协议范围内。",
        PALE_GREEN,
    )
    add_label_paragraph(
        doc,
        "最重要的表达原则",
        "不要把结构化执行写成端到端 VLA、LoRA 后训练或真实机械臂结果；不要把 40 条低摩擦配对结果写成统计显著。",
        PALE_GOLD,
    )
    add_paragraph(doc, "版本依据：mujoco_research_summary_v1；指南生成日期：2026-07-22。",
                  size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=0)
    doc.add_page_break()


def add_timing_route(doc):
    add_heading(doc, "演讲路线：按网页顺序，但不被网页顺序束缚", 1)
    add_paragraph(doc, "先切换到网页中文界面。讲解时以数据图为主、视频为辅；下表是一条 9 分钟的推荐路线。", after=7)
    add_data_table(
        doc,
        ["时间", "网页位置", "你要完成的事", "讲完后的听众认知"],
        [
            ("0:00-0:40", "顶部指标 + 研究问题", "先给可防守的结论和研究范围", "知道项目不是端到端 VLA"),
            ("0:40-2:10", "标准搬运对比", "解释 0/100 的严格成功率和接触控制缺口", "明白表征不等于抓取控制"),
            ("2:10-3:40", "低摩擦分任务渐变", "用四任务 30/40 到 36/40 展示可靠性趋势", "看到不只有绝对结果"),
            ("3:40-5:00", "语言与干扰验证", "分开讲事后修复和独立句法", "理解语言结论的边界"),
            ("5:00-5:45", "数据效率与资源", "说明冻结模型、小样本和算力口径", "理解轻量化的真实含义"),
            ("5:45-6:30", "视频证据", "播放 1-2 段成功视频，必要时补失败", "把指标与可视过程对应起来"),
            ("6:30-8:00", "方法表 + 融合区", "总结方法迭代、局限和下一步", "知道结果可复现但仍有限制"),
        ],
        [1100, 1850, 3500, 2910],
        font_size=9,
    )
    add_label_paragraph(doc, "现场操作", "打开本地 HTML 后保持中文；先滚到“研究问题与可写结论”，不要从顶部的视频开始。切换图表时用鼠标指向当前数据行，不要同时谈两个图。", PALE_GOLD)

    add_heading(doc, "0. 开场 40 秒：先建立问题、方法和边界", 1)
    add_script(doc, "我的问题是在有限示范和有限算力下，视觉语言表征能否帮助 WidowX 完成桌面搬运。最终结果不是训练出端到端的大型 VLA，而是验证了一个分层结论：冻结 CLIP 可以稳定完成任务级语义选择；真正保证抓取与放置的是结构化接触执行。后面我会用标准任务、低摩擦压力、语言测试和视频证据分别说明这四层结论。")
    add_paragraph(doc, "说完后停半秒，再滚到“标准搬运任务”。这样听众先有判断框架，后面的 0% 与 100% 就不会被误解成简单的模型排行榜。")


def add_canonical_section(doc):
    add_heading(doc, "1. 如何讲“标准搬运任务：学习控制与分层方案对比”", 1)
    add_paragraph(doc, "这一区是全场最容易被追问的地方。不要把 0% 与 100%当作单纯的优劣排名，而要把它解释为一次明确的控制结构诊断。", after=7)
    add_label_paragraph(doc, "严格成功口径", "语义意图正确、物体抬升至少 0.06 m 且与 TCP 持续接近、最终放入目标区域。只接近目标或偶然碰到目标都不算成功。", PALE_GOLD)
    add_data_table(
        doc,
        ["方法类型", "标准任务", "平均最终目标距离", "应讲出的结论"],
        [
            ("线性 BC / Trajectory-kNN / 对象-语言动作头", "0/20", "约 0.225-0.237 m", "普通回归或检索没有形成有效抓取"),
            ("冻结 CLIP 连续动作头", "0/20", "0.2382 m", "图文表征本身不能解决接触控制"),
            ("CLIP 语义 + 结构化执行", "20/20", "0.0100 m", "语义与接触执行解耦后可稳定搬运"),
            ("采用：加闭词表规范化", "20/20", "0.0100 m", "保留结构化执行，修复词表内同义表达"),
        ],
        [2620, 1170, 1620, 3950],
        font_size=9.2,
    )
    add_script(doc, "这里的 0/20 不是说 CLIP 没有语义能力，而是说明把冻结图文特征直接回归成连续关节或夹爪动作，仍然无法处理接触、夹紧和抬升。把语义选择与结构化接触执行解耦后，标准协议达到 20/20。这是控制分工的证据，不是端到端 VLA 已经解决抓取的证据。")
    add_heading(doc, "面对“数据太绝对”的解释方式", 2)
    add_paragraph(doc, "网页中 0% 与 100%是严格门槛下的饱和结果，保留它是为了突出“是否形成有效抓取”的结构性差异。演讲时要立即把视线转到同一区块的平均最终目标距离，以及后面的低摩擦四任务渐变数据。这样既保留了清晰的负例，也避免把全篇工作讲成只有一个二元指标。")
    add_label_paragraph(doc, "不要这样说", "“我的方法比所有方法强 100%。” 这会把小规模固定协议的门槛结果错误外推成通用性能排序。", PALE_RED)
    add_label_paragraph(doc, "可以这样说", "“在这 20 个标准留出 episode 的严格抓取门槛下，连续动作回归没有形成有效抓取；采用的分层执行达到饱和，因此我们继续用更困难的低摩擦协议观察渐变差异。”", PALE_GREEN)


def add_fusion_section(doc):
    add_heading(doc, "2. 用低摩擦压力测试给出“渐变性”对比", 1)
    add_paragraph(doc, "这是回应“0%/100%过于绝对”的主证据。所有任务使用相同初始状态、任务、hard 物体复杂度和冻结 CLIP adapter，只改变执行器：标准语义执行器与带接触/接近反馈的融合执行器。每个任务各有 10 个配对 seed。", after=7)
    add_data_table(
        doc,
        ["任务", "标准", "融合", "变化", "平均目标距离（m）"],
        [
            ("蓝方块 -> 蓝盘", "8/10", "10/10", "+20 pp", "0.0481 -> 0.0138"),
            ("蓝方块 -> 红盘", "8/10", "9/10", "+10 pp", "0.0690 -> 0.0259"),
            ("红方块 -> 红盘", "7/10", "10/10", "+30 pp", "0.0575 -> 0.0145"),
            ("最左方块 -> 碗", "7/10", "7/10", "+0 pp", "0.0527 -> 0.0467"),
            ("合计", "30/40", "36/40", "+15 pp", "0.0568 -> 0.0252"),
        ],
        [2600, 1050, 1050, 1250, 3410],
        header_fill=TEAL,
        font_size=9.5,
    )
    add_script(doc, "在标准摩擦下两个执行器都达到 20/20，因此无法区分可靠性。低摩擦后才出现梯度：四个任务总体从 30/40 提升到 36/40，平均最终距离从 5.68 厘米降到 2.52 厘米。差异并不均匀，最左方块到碗没有成功率提升，所以这个结果是物理鲁棒性改善趋势，而不是对所有任务都必然有效的结论。")
    add_heading(doc, "统计边界必须主动讲", 2)
    add_paragraph(doc, "40 条配对轨迹中，失败转成功 7 条、成功转失败 1 条，精确 McNemar 双侧 p=0.0703。这个值支持“改善趋势”，但不满足常用的 p<0.05 显著性阈值。主动说出这一点会比回避它更有说服力。")
    add_label_paragraph(doc, "融合器的真实构成", "它从第一次抓取起就采用更紧、更长的夹爪保持，并在失物或未放置时最多重规划一次。因此收益不能单独归因于恢复分支；恢复分支的独立消融仍未完成。", PALE_GOLD)
    add_label_paragraph(doc, "视频中的融合诊断", "网页最后一个融合视频最终满足原严格成功，但最终运输保持代理为 false。演示时只能称为“恢复过程诊断”，不要称为“全程稳定抓取的成功案例”。", PALE_RED)


def add_language_resource_section(doc):
    add_heading(doc, "3. 如何讲语言与干扰验证", 1)
    add_paragraph(doc, "这一节的目标不是证明开放词汇泛化，而是区分：原始改写中的错误能否修复，以及在固定闭词表下能否理解新的句式。", after=7)
    add_data_table(
        doc,
        ["测试", "协议", "采用方案结果", "该结果能证明什么"],
        [
            ("原始改写集", "60 episode；事后修复", "60/60", "闭词表别名规范化修复原协议中的颜色、形状与目标区同义词"),
            ("全物体干扰", "20 episode", "20/20", "在当前物体集合中可排除干扰对象"),
            ("独立句法", "8 条新句子 × 4 任务 × 5 seed", "40/40", "固定闭词表内的新句式可以被正确解释"),
        ],
        [1500, 2170, 1250, 4440],
        font_size=9.3,
    )
    add_script(doc, "原始改写集从基础版本的 51/60，到采用规范化后的 60/60，是对原协议错误的事后修复，不是独立 OOD 证据。独立证据是后面的 8 条新完整句子：在固定闭词表内达到 40/40。因此我只声称它能处理词表内的新句式，不声称开放词汇语言泛化。")
    add_label_paragraph(doc, "负例也要保留", "训练时加入语义改写增强后，原始改写集变为 48/60、全物体干扰变为 16/20，反而退化。它说明当前小规模设置下，手工别名规范化比该增强策略更可靠。", PALE_GOLD)

    add_heading(doc, "4. 如何讲数据效率与资源口径", 1)
    add_data_table(
        doc,
        ["项目", "网页数值", "演讲时的正确解释"],
        [
            ("冻结图文编码器", "151,277,313 参数", "大部分表征被冻结，不是在本机从零训练大模型"),
            ("可训练意图 adapter", "4,100 参数", "轻量化发生在四类任务意图适配层"),
            ("训练与显存", "8.44 秒；678.7 MB 峰值", "说明该小规模意图适配可在有限资源下运行"),
            ("示范", "79 条成功示范", "当前任务与固定留出集上的数据规模，不代表复杂机器人数据上限"),
            ("每类示范预算", "20 / 40 / 79 样本均为 20/20", "固定留出集已饱和，不能把它解读为完整的数据缩放曲线"),
        ],
        [2000, 1850, 5510],
        font_size=9.3,
    )
    add_script(doc, "轻量化在这里的含义很具体：151M 参数的 CLIP 编码器冻结，只训练 4,100 个意图 adapter 参数，用 79 条示范完成四类任务选择。网页中的示范预算都达到 20/20，说明这个固定留出集已经饱和；因此它证明的是当前任务上的低资源可行性，而不是一般意义上的数据缩放规律。")


def add_video_section(doc):
    add_heading(doc, "5. 怎样使用“视频证据”而不让视频替代数据", 1)
    add_paragraph(doc, "视频的工作是把定量结果变成可检查的过程证据。建议最多播放两段成功视频，总时长控制在 25 秒以内；失败和融合诊断只在需要解释设计动机时播放。", after=7)
    add_data_table(
        doc,
        ["视频", "建议时机", "你要指出什么", "不可据此声称什么"],
        [
            ("蓝方块 -> 蓝盘", "标准任务后，播放 8-12 秒", "颜色对象与目标盘匹配、抓取、抬升、放置", "单条视频不等于整体成功率"),
            ("红方块 -> 红盘", "语言节后，播放 8-12 秒", "不同颜色任务也遵循同一结构化执行流程", "不能说是开放词汇泛化"),
            ("最左方块 -> 碗", "可替换第二段成功视频", "空间关系“最左”的语义选择", "不能证明所有空间关系都支持"),
            ("连续动作头失败", "解释 0/20 时，播放约 6 秒", "没有形成有效抓取，与严格门槛一致", "失败视频不是所有对照方法的完整代表"),
            ("融合恢复诊断", "仅在被问到接触反馈时", "低摩擦脱落后的有限恢复过程", "不能称为全程稳定持物"),
        ],
        [1500, 2050, 3100, 2710],
        font_size=8.8,
    )
    add_script(doc, "我用视频只验证两个观察点：第一，蓝方块能够被放入蓝盘；第二，最左方块到碗体现空间关系选择。视频不是成功率来源，成功率来自固定 seed 下的批量评测；视频用于让大家检查指标对应的实际运动过程。")
    add_heading(doc, "网页现场操作顺序", 2)
    add_paragraph(doc, "1. 先把浏览器缩放固定在 100%，选择中文；2. 讲数据时停在图表标题和数值上；3. 滚到“视频证据”后只点击一个视频；4. 物体释放并进入盘/碗后立即暂停或让视频结束；5. 回到“未解决问题”结束，而不是让听众连续观看六个视频。")
    add_label_paragraph(doc, "本地页面", "C:\\Users\\Administrator\\Desktop\\final project-robot ic\\vla_robot_grasping\\docs\\mujoco_research_summary_zh_en.html。答辩前打开一次并确认六个视频可以加载。", LIGHT_BLUE)


def add_questions_and_close(doc):
    add_heading(doc, "6. 答辩追问：可直接使用的回答", 1)
    add_question(doc, "为什么标准对比里很多方法是 0%，这是否让数据过于绝对？", "这些是严格抓取门槛下的结果，目的是诊断是否形成有效接触控制，而不是声称通用排名。为避免只给出二元结果，我在低摩擦协议中报告了四个任务、40 条配对 seed、成功率和最终距离的渐变数据。")
    add_question(doc, "你的方法是 VLA 或 LoRA 后训练吗？", "不是。当前版本冻结 CLIP，只训练四类任务意图的线性 adapter；抓取与放置由结构化 waypoint 执行器完成。它是轻量视觉语言语义适配加结构化控制的 MuJoCo 原型，不是端到端 VLA 或 LoRA。")
    add_question(doc, "60/60 的语言结果能证明语言泛化吗？", "只能证明原始改写协议在闭词表规范化后被修复。独立证据是 8 条新完整句子上的 40/40，它也仍局限于固定颜色、形状、目标区域和左右关系词表。")
    add_question(doc, "30/40 到 36/40 是否统计显著？", "不是。配对精确 McNemar 检验的双侧 p=0.0703，因此我把它称为低摩擦下的改善趋势，而不是统计显著结论。下一步需要更多 seed、扰动等级和独立训练/测试划分。")
    add_question(doc, "接触反馈恢复究竟带来了多少提升？", "当前融合器从第一次抓取起就改变了夹爪保持，并在必要时允许一次重规划，所以整体收益不能单独归因于恢复分支。它的价值是提出了一个可测的执行层补强方向，独立消融仍是下一步工作。")
    add_question(doc, "这能迁移到真实 WidowX 吗？", "当前没有真实机械臂验证。所有结论严格限定在 MuJoCo 的四类桌面任务和固定评测协议；真实系统还需要传感、控制频率、安全和 sim-to-real 的独立实验。")

    add_heading(doc, "7. 最后 30 秒：用边界收束，而不是夸大结论", 1)
    add_script(doc, "最后总结：本项目的贡献不是宣布端到端 VLA 已经学会抓取，而是在 MuJoCo 的 WidowX 桌面任务中明确分离了两件事。冻结视觉语言表征可以用极少可训练参数完成任务级语义选择；接触可靠性仍需要结构化执行与反馈。标准任务给出清晰的控制差异，低摩擦压力测试给出 30/40 到 36/40 的改善趋势，语言实验明确限定在闭词表范围。下一步将把抓取和抬升子策略改为可学习的闭环模块，并扩大独立 OOD 与扰动评测。")
    add_label_paragraph(doc, "结束前检查", "页面中文已打开；能播放蓝方块到蓝盘与最左方块到碗；记住 p=0.0703 不是显著；记住“MuJoCo-only、非端到端 VLA、闭词表”。", PALE_GREEN)

    doc.add_page_break()
    add_heading(doc, "附：网页事实卡", 1)
    add_data_table(
        doc,
        ["事实", "数值 / 表述"],
        [
            ("主方法", "冻结 CLIP 意图选择 + 闭词表规范化 + 结构化接触执行"),
            ("标准任务", "4 类任务，20/20 严格成功；平均目标距离 0.0100 m"),
            ("连续 CLIP 动作头", "0/20；平均目标距离 0.2382 m"),
            ("低摩擦压力", "标准 30/40 -> 融合 36/40；p=0.0703，改善趋势但不显著"),
            ("独立句法", "固定闭词表内 40/40；不代表开放词汇"),
            ("训练口径", "151,277,313 冻结编码器参数；4,100 可训练 adapter 参数；79 条示范"),
            ("外推边界", "MuJoCo-only；没有真实机器人、端到端 VLA、LoRA 或 RL 后训练结论"),
        ],
        [2360, 7000],
        font_size=9.5,
    )
    add_paragraph(doc, "数据来源：当前 HTML 内嵌的 mujoco_research_summary_v1 数据，以及对应的接触融合评测报告。", size=9, color=MUTED, before=4, after=0)


def build_document():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_timing_route(doc)
    add_canonical_section(doc)
    add_fusion_section(doc)
    add_language_resource_section(doc)
    add_video_section(doc)
    add_questions_and_close(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
